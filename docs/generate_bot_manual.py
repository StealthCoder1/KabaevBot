from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT_DIR = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT_DIR / "docs"
SCREENSHOTS_DIR = DOCS_DIR / "screenshots"
OUTPUT_DOCX = DOCS_DIR / "APHotlotsbot_manual.docx"


def _font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = []
    if bold:
        candidates.extend(
            [
                Path("C:/Windows/Fonts/arialbd.ttf"),
                Path("C:/Windows/Fonts/segoeuib.ttf"),
            ]
        )
    else:
        candidates.extend(
            [
                Path("C:/Windows/Fonts/arial.ttf"),
                Path("C:/Windows/Fonts/segoeui.ttf"),
            ]
        )

    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def _draw_phone_screen(
    image_path: Path,
    *,
    title: str,
    subtitle: str,
    blocks: list[str],
    footer: str | None = None,
) -> None:
    width, height = 1080, 1920
    image = Image.new("RGB", (width, height), "#dfe7ee")
    draw = ImageDraw.Draw(image)

    draw.rectangle((0, 0, width, 260), fill="#2a5d84")
    draw.text((60, 54), "Telegram", fill="white", font=_font(44, bold=True))
    draw.text((60, 130), "@APHotlotsbot", fill="#dceefe", font=_font(36))

    phone_left, phone_top = 70, 280
    phone_right, phone_bottom = width - 70, height - 110
    draw.rounded_rectangle(
        (phone_left, phone_top, phone_right, phone_bottom),
        radius=42,
        fill="#f9fbfd",
        outline="#90a4b6",
        width=4,
    )

    title_x = phone_left + 48
    title_y = phone_top + 48
    draw.text((title_x, title_y), title, fill="#1b2d3f", font=_font(40, bold=True))
    draw.multiline_text(
        (title_x, title_y + 72),
        subtitle,
        fill="#36506a",
        font=_font(28),
        spacing=8,
    )

    block_top = title_y + 180
    for block_text in blocks:
        block_height = 108
        draw.rounded_rectangle(
            (title_x, block_top, phone_right - 48, block_top + block_height),
            radius=24,
            fill="#eaf3fb",
            outline="#b7d2e8",
            width=2,
        )
        draw.text((title_x + 26, block_top + 32), block_text, fill="#173756", font=_font(28))
        block_top += block_height + 22

    if footer:
        draw.multiline_text(
            (title_x, phone_bottom - 180),
            footer,
            fill="#5f6f7f",
            font=_font(24),
            spacing=6,
        )

    image_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(image_path)


def _generate_screenshots() -> list[Path]:
    shots: list[Path] = []

    main_menu = SCREENSHOTS_DIR / "01_main_menu.png"
    _draw_phone_screen(
        main_menu,
        title="Главное меню",
        subtitle="После /start пользователь\nвидит основные разделы:",
        blocks=[
            "Максимальная выгода",
            "Подборка авто",
            "Подборка мото",
            "Авто в пути",
            "Гарантии",
            "Быстро о главном",
            "Связаться с менеджером",
        ],
        footer="Скриншот-схема интерфейса",
    )
    shots.append(main_menu)

    admin_menu = SCREENSHOTS_DIR / "02_admin_menu.png"
    _draw_phone_screen(
        admin_menu,
        title="Админ-панель",
        subtitle="Открывается командой /admin\nдля ID из ADMIN_TG_ID:",
        blocks=[
            "Статистика пользователей",
            "Последние лиды",
            "Канал авто в пути",
            "Канал лидов",
        ],
        footer="Если доступа нет, бот отвечает: \"Доступ запрещен\"",
    )
    shots.append(admin_menu)

    auto_channel = SCREENSHOTS_DIR / "03_auto_channel_setup.png"
    _draw_phone_screen(
        auto_channel,
        title="Канал Авто в пути",
        subtitle="Настройка в админке:",
        blocks=[
            "1) Нажмите: Канал авто в пути",
            "2) Отправьте chat_id, например:",
            "   -1003706573371",
            "3) Получите подтверждение",
        ],
        footer="Важно: при смене ID бот очищает сохраненные\nпосты 'Авто в пути' в базе.",
    )
    shots.append(auto_channel)

    leads_channel = SCREENSHOTS_DIR / "04_leads_channel_setup.png"
    _draw_phone_screen(
        leads_channel,
        title="Канал лидов",
        subtitle="Настройка в админке:",
        blocks=[
            "1) Нажмите: Канал лидов",
            "2) Отправьте chat_id канала",
            "3) Лиды начнут приходить туда",
        ],
        footer="Дополнительно лиды приходят в личные сообщения\nадминам из ADMIN_TG_ID.",
    )
    shots.append(leads_channel)

    bot_admin = SCREENSHOTS_DIR / "05_bot_admin_rights.png"
    _draw_phone_screen(
        bot_admin,
        title="Как выдать права боту",
        subtitle="Telegram канал/группа:",
        blocks=[
            "1) Управление каналом",
            "2) Администраторы",
            "3) Добавить @APHotlotsbot",
            "4) Сохранить права",
        ],
        footer="Рекомендуется разрешить публикацию сообщений.\nДля групп также отключите Privacy Mode в BotFather.",
    )
    shots.append(bot_admin)

    return shots


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    if heading.runs:
        heading.runs[0].font.name = "Calibri"


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(item, style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)


def _add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(item, style="List Number")
        paragraph.paragraph_format.space_after = Pt(3)


def _build_docx(screenshots: list[Path]) -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    title = doc.add_heading("Инструкция по боту @APHotlotsbot", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    created_at = doc.add_paragraph(
        f"Версия документа: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    )
    created_at.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        "Документ описывает полный функционал бота, работу с лидами, "
        "настройку каналов и выдачу прав админа."
    )

    _add_heading(doc, "1. Полный функционал бота")
    _add_bullets(
        doc,
        [
            "Главное меню: максимальная выгода, подбор авто, подбор мото, авто в пути, гарантии, быстро о главном, связь с менеджером.",
            "Подбор авто: выбор бюджета, выбор модели, карточка модели, кнопка «Хочу это авто», создание лида.",
            "Подбор мото: выбор класса (спорт/стрит/круизер), выбор модели, кнопка «Хочу этот мото», создание лида.",
            "Максимальная выгода: показ случайного лота из JSON-каталога с возможностью листать следующий.",
            "Авто в пути: бот копирует накопленные посты из настроенного канала/группы пользователю по кнопке.",
            "Гарантии и раздел «Быстро о главном»: информационные блоки с кнопками перехода к подборкам и менеджеру.",
            "Сбор контактов: кнопка отправки контакта или ввод в формате «Имя +79991234567».",
            "Лиды: сохраняются в БД и отправляются в канал лидов + администраторам.",
            "Админка: статистика пользователей, последние лиды, настройка канала «Авто в пути», настройка канала лидов.",
            "Фоновая функция: раз в 24 часа отправляется вопрос по бренду всем пользователям из таблицы users.",
            "Рассылка постов из источников: новые посты в источнике отправляются всем пользователям; поддерживаются текст, фото, видео и медиа-группы.",
        ],
    )

    _add_heading(doc, "2. Как добавить канал для «Авто в пути» (подробно)")
    doc.add_paragraph(
        "Ниже точный порядок действий: куда зайти, что нажать и как проверить, что настройка применена."
    )

    _add_heading(doc, "2.1 Добавьте бота в канал/группу", level=2)
    _add_numbers(
        doc,
        [
            "Откройте Telegram и зайдите в нужный канал (или группу), который будет источником постов «Авто в пути».",
            "Нажмите на название канала вверху экрана.",
            "Откройте «Управление каналом» (или «Управление группой»).",
            "Зайдите в раздел «Администраторы».",
            "Нажмите «Добавить администратора».",
            "Найдите и выберите @APHotlotsbot.",
            "Сохраните права администратора для бота.",
        ],
    )
    doc.add_paragraph(
        "Рекомендуемые права: публикация сообщений. Для группы дополнительно отключите Privacy Mode через BotFather, "
        "чтобы бот видел обычные сообщения в группе."
    )

    _add_heading(doc, "2.2 Получите chat_id через @myidbot", level=2)
    _add_numbers(
        doc,
        [
            "Откройте ваш канал/группу и отправьте любое тестовое сообщение (например: «test id»).",
            "Нажмите «Переслать» на этом сообщении.",
            "Выберите чат @myidbot и отправьте пересланное сообщение ему.",
            "В ответе @myidbot найдите ID чата/канала. Нужное значение обычно начинается с -100.",
            "Скопируйте это число полностью (пример: -1003706573371).",
        ],
    )
    doc.add_paragraph(
        "Если ID пришел без «-100», проверьте, что переслали именно сообщение из нужного канала/группы."
    )

    _add_heading(doc, "2.3 Внесите chat_id в админке бота", level=2)
    _add_numbers(
        doc,
        [
            "Откройте диалог с @APHotlotsbot.",
            "Отправьте команду /admin.",
            "Нажмите кнопку «Канал авто в пути».",
            "Отправьте в чат скопированный chat_id одним сообщением (только число, без текста).",
            "Дождитесь сообщения-подтверждения от бота, что канал обновлен.",
        ],
    )

    _add_heading(doc, "2.4 Проверьте, что всё работает", level=2)
    _add_numbers(
        doc,
        [
            "Опубликуйте тестовый пост в канале/группе «Авто в пути».",
            "Откройте бота как обычный пользователь и нажмите кнопку «Авто в пути» в главном меню.",
            "Убедитесь, что бот отправил этот пост пользователю.",
        ],
    )
    doc.add_paragraph(
        "Важно: если вы меняете chat_id «Авто в пути», бот очищает старые сохраненные посты этой рубрики в базе данных."
    )

    _add_heading(doc, "3. Как настроить канал для лидов")
    _add_numbers(
        doc,
        [
            "Добавьте бота @APHotlotsbot в канал/группу для лидов как администратора.",
            "Откройте /admin и нажмите «Канал лидов».",
            "Отправьте chat_id канала/группы (числом, например -100...).",
            "После подтверждения все новые лиды будут отправляться в этот канал.",
        ],
    )
    doc.add_paragraph(
        "Параллельно лиды продолжают отправляться администраторам, чьи ID указаны в ADMIN_TG_ID."
    )

    _add_heading(doc, "4. Как сделать бота админом")
    _add_numbers(
        doc,
        [
            "Откройте нужный канал или группу в Telegram.",
            "Перейдите в «Управление» -> «Администраторы».",
            "Нажмите «Добавить администратора» и выберите @APHotlotsbot.",
            "Выдайте права (минимум: публикация сообщений в канале).",
            "Сохраните изменения.",
        ],
    )
    doc.add_paragraph(
        "Если используете группу как источник «Авто в пути», рекомендуется отключить Privacy Mode у бота через BotFather, чтобы бот видел обычные сообщения группы."
    )

    _add_heading(doc, "5. Команды и кнопки администратора")
    _add_bullets(
        doc,
        [
            "/admin — открыть админ-панель (только для ADMIN_TG_ID).",
            "/stat или /stats — статистика по пользователям.",
            "Кнопка «Статистика пользователей» — то же, что /stat.",
            "Кнопка «Последние лиды» — вывод последних 10 лидов.",
            "Кнопка «Канал авто в пути» — смена источника постов для раздела «Авто в пути».",
            "Кнопка «Канал лидов» — канал, куда отправляются новые лиды.",
        ],
    )

    _add_heading(doc, "6. Примеры экранов (скриншоты-схемы)")
    for idx, shot in enumerate(screenshots, start=1):
        doc.add_paragraph(f"Скриншот {idx}: {shot.stem.replace('_', ' ')}")
        doc.add_picture(str(shot), width=Inches(5.7))

    _add_heading(doc, "7. Быстрая проверка после настройки")
    _add_numbers(
        doc,
        [
            "Выполните /start и проверьте, что открывается главное меню.",
            "Выполните /admin с админ-аккаунта и проверьте кнопки админки.",
            "Отправьте тестовый пост в канал «Авто в пути» и проверьте выдачу по кнопке «Авто в пути».",
            "Отправьте тестовый лид (контакт) и убедитесь, что он пришел в канал лидов.",
        ],
    )

    doc.save(OUTPUT_DOCX)


def main() -> None:
    screenshots = _generate_screenshots()
    _build_docx(screenshots)
    print(f"Created: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
